"""Text extraction from PDF, DOCX, and TXT files."""
import io
from utils import get_file_extension, clean_text


class TextExtractionError(Exception):
    """Error during text extraction."""
    pass


def extract_text(file_path: str = None, file_content: bytes = None, filename: str = None) -> str:
    """Extract text from a file.

    Args:
        file_path: Path to local file (optional)
        file_content: Raw file content as bytes (optional)
        filename: Original filename to determine type (required if using file_content)

    Returns:
        Extracted text as string
    """
    if file_path:
        extension = get_file_extension(file_path)
        with open(file_path, 'rb') as f:
            content = f.read()
    elif file_content and filename:
        extension = get_file_extension(filename)
        content = file_content
    else:
        raise TextExtractionError("Must provide either file_path or (file_content and filename)")

    extractors = {
        'pdf': extract_from_pdf,
        'docx': extract_from_docx,
        'doc': extract_from_docx,  # Try docx parser for .doc
        'txt': extract_from_txt
    }

    extractor = extractors.get(extension)
    if not extractor:
        raise TextExtractionError(f"Unsupported file type: {extension}. Supported: pdf, docx, txt")

    text = extractor(content)
    return clean_text(text)


def extract_from_pdf(content: bytes) -> str:
    """Extract text from PDF content."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(content))
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        if not text_parts:
            raise TextExtractionError("No text could be extracted from PDF. It may be scanned/image-based.")

        return '\n\n'.join(text_parts)

    except ImportError:
        raise TextExtractionError("PyPDF2 library not available")
    except Exception as e:
        raise TextExtractionError(f"Error extracting PDF text: {str(e)}")


def extract_from_docx(content: bytes) -> str:
    """Extract text from DOCX content."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(' | '.join(row_text))

        if not text_parts:
            raise TextExtractionError("No text could be extracted from document.")

        return '\n\n'.join(text_parts)

    except ImportError:
        raise TextExtractionError("python-docx library not available")
    except Exception as e:
        raise TextExtractionError(f"Error extracting DOCX text: {str(e)}")


def extract_from_txt(content: bytes) -> str:
    """Extract text from TXT content."""
    try:
        # Try UTF-8 first, then fall back to latin-1
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            return content.decode('latin-1')
    except Exception as e:
        raise TextExtractionError(f"Error reading text file: {str(e)}")
